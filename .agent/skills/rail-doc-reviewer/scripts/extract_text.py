"""
extract_text.py — Batch text extraction for Downer QTMP project documents.

Converts PDF, DOCX, and XLSX files to plain-text markdown for agent consumption.
Outputs go to: .agent/skills/rail-doc-reviewer/resources/raw_text/

Usage:
    python .agent/skills/rail-doc-reviewer/scripts/extract_text.py [--force]

Options:
    --force   Re-extract all documents even if output already exists.
"""

import os
import sys
import re
import hashlib
from pathlib import Path

# ── Paths ───────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
RESOURCES_DIR = SKILL_DIR / "resources"
RAW_TEXT_DIR = RESOURCES_DIR / "raw_text"
PROJECT_ROOT = SKILL_DIR.parent.parent.parent  # Up from .agent/skills/rail-doc-reviewer/scripts/

# ── Ensure output dirs exist ────────────────────────────────────────────────
RAW_TEXT_DIR.mkdir(parents=True, exist_ok=True)


def sanitise_filename(name: str) -> str:
    """Convert a document filename to a safe, readable stem for the output file."""
    stem = Path(name).stem
    # Replace spaces, dots, hyphens, and underscores with underscores; collapse multiples
    safe = re.sub(r"[\s\.\-]+", "_", stem)
    safe = re.sub(r"_+", "_", safe).strip("_")
    # Truncate to a reasonable length
    if len(safe) > 80:
        safe = safe[:80]
    return safe


def extract_pdf(filepath: Path) -> str:
    """Extract all text from a PDF using pdfplumber."""
    import pdfplumber

    lines = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                lines.append(f"--- PAGE {i} ---")
                lines.append(text)
            else:
                lines.append(f"--- PAGE {i} --- [NO TEXT EXTRACTED]")

            # Also try to extract tables
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                lines.append(f"\n[TABLE {t_idx + 1} on page {i}]")
                for row in table:
                    cleaned = [str(cell).replace("\n", " ").strip() if cell else "" for cell in row]
                    lines.append(" | ".join(cleaned))

    return "\n".join(lines)


def extract_docx(filepath: Path) -> str:
    """Extract all text from a DOCX using python-docx."""
    from docx import Document

    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        # Attempt to preserve heading structure
        if "Heading 1" in style_name:
            lines.append(f"\n# {text}")
        elif "Heading 2" in style_name:
            lines.append(f"\n## {text}")
        elif "Heading 3" in style_name:
            lines.append(f"\n### {text}")
        elif "Heading 4" in style_name:
            lines.append(f"\n#### {text}")
        elif "List" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Extract tables
    for t_idx, table in enumerate(doc.tables):
        lines.append(f"\n[TABLE {t_idx + 1}]")
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_xlsx(filepath: Path) -> str:
    """Extract all data from an XLSX using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(filepath, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"\n=== SHEET: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cells):  # Skip entirely blank rows
                lines.append(" | ".join(cells))

    return "\n".join(lines)


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
}


def find_documents(root: Path) -> list[Path]:
    """Find all extractable documents in the project root (non-recursive into .agent)."""
    docs = []
    for item in sorted(root.rglob("*")):
        # Skip the .agent directory itself
        if ".agent" in item.parts:
            continue
        if item.is_file() and item.suffix.lower() in EXTRACTORS:
            docs.append(item)
    return docs


def main():
    force = "--force" in sys.argv

    docs = find_documents(PROJECT_ROOT)
    print(f"Found {len(docs)} documents to process.\n")

    results = {"extracted": 0, "skipped": 0, "failed": 0}

    for doc in docs:
        safe_name = sanitise_filename(doc.name)
        out_path = RAW_TEXT_DIR / f"{safe_name}.txt"

        if out_path.exists() and not force:
            print(f"  SKIP (exists): {doc.name}")
            results["skipped"] += 1
            continue

        ext = doc.suffix.lower()
        extractor = EXTRACTORS.get(ext)
        if not extractor:
            print(f"  SKIP (unsupported): {doc.name}")
            results["skipped"] += 1
            continue

        try:
            print(f"  EXTRACTING: {doc.name} ...")
            text = extractor(doc)

            # Write header + extracted text
            header = (
                f"SOURCE FILE: {doc.name}\n"
                f"FULL PATH: {doc}\n"
                f"EXTENSION: {ext}\n"
                f"SIZE: {doc.stat().st_size} bytes\n"
                f"{'=' * 60}\n\n"
            )
            out_path.write_text(header + text, encoding="utf-8")

            word_count = len(text.split())
            print(f"    -> {out_path.name} ({word_count} words)")
            results["extracted"] += 1

        except Exception as e:
            print(f"  FAILED: {doc.name} — {e}")
            # Write a stub so we know it was attempted
            out_path.write_text(
                f"SOURCE FILE: {doc.name}\nFULL PATH: {doc}\nEXTRACTION FAILED: {e}\n",
                encoding="utf-8",
            )
            results["failed"] += 1

    print(f"\nDone. Extracted: {results['extracted']}, Skipped: {results['skipped']}, Failed: {results['failed']}")
    print(f"Output directory: {RAW_TEXT_DIR}")


if __name__ == "__main__":
    main()
